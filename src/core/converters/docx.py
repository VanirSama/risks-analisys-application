from src.core.converters.base import Converter
from src.models.risk_map import RiskMapFile
from src.utils.resources import RESOURCE_LOADER

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path


class RiskMapToDocxConverter(Converter[RiskMapFile]):
    FILTER_STRING = "Документы MS Word (*.docx)"
    OUTPUT_FORMAT = "docx"

    def __init__(self, risk_map: RiskMapFile):
        super().__init__(risk_map)
        self._file: RiskMapFile = risk_map

        self._document = Document()
        self._setup_document()
        self._set_document_metadata()

    def convert(self, output_path: Path | str) -> None:
        self._add_logo()
        self._add_approval_section()
        self._add_title()
        self._add_main_info()
        self._add_summary_table()
        self._add_risk_table()
        self._add_methods()
        self._document.save(output_path)

    def _set_document_metadata(self) -> None:
        core_properties = self._document.core_properties
        core_properties.author = self._METADATA.get("author", "")
        core_properties.comments = self._METADATA.get("comments", "")

    def _setup_document(self) -> None:
        self._document.styles['Normal'].font.name = 'Times New Roman'

        section = self._document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    def _add_logo(self) -> None:
        self._document.add_picture(RESOURCE_LOADER.get("LOGO_COLORED", ""), width=Inches(1))
        logo_paragraph = self._document.paragraphs[-1]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = logo_paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)

    def _add_approval_section(self) -> None:
        table = self._document.add_table(rows=3, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Inches(3.5)

        cell_00 = table.rows[0].cells[0]
        cell_00.text = "Утверждаю"
        cell_00.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell_10 = table.rows[1].cells[0]
        cell_10.text = "Председатель комиссии"
        cell_10.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell_20 = table.rows[2].cells[0]
        chairman_paragraph = cell_20.paragraphs[0]
        chairman_paragraph.add_run("______________________________________").font.size = Pt(10)

        chairman = self._file.chairman
        if not chairman: chairman = "Фамилия И.О."
        chairman_paragraph.add_run(chairman).font.size = Pt(10)
        chairman_paragraph.add_run("\n______________________________20____г.").font.size = Pt(10)
        chairman_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_title(self) -> None:
        title_table = self._document.add_table(rows=1, cols=1)
        title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        title_table.autofit = False

        cell_00 = title_table.rows[0].cells[0]
        map_paragraph = cell_00.paragraphs[0]
        map_paragraph.add_run(f"Карта оценки риска №{self._file.map_no if self._file.map_no else '___'}").font.bold = True
        map_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._document.add_paragraph()

    def _add_main_info(self) -> None:
        profession_paragraph = self._document.add_paragraph()
        profession_paragraph.add_run("1. Название профессии\n")
        profession_paragraph.add_run(self._file.profession if self._file.profession else "")
        profession_paragraph.runs[0].font.bold = True

        department_paragraph = self._document.add_paragraph()
        department_paragraph.add_run(f"2. Наименование структурного подразделения\n")
        department_paragraph.add_run(self._file.department if self._file.department else "")
        department_paragraph.runs[0].font.bold = True

        description_paragraph = self._document.add_paragraph()
        description_paragraph.add_run("3. Краткое описание выполняемой работы\n")
        description_paragraph.add_run(self._file.description if self._file.description else "")
        description_paragraph.runs[0].font.bold = True

        tools_and_materials_paragraph = self._document.add_paragraph()
        tools_and_materials_paragraph.add_run("4. Используемое оборудование, материалы, сырье\n")
        tools_and_materials_paragraph.add_run(self._file.tools_and_materials if self._file.tools_and_materials else "")
        tools_and_materials_paragraph.runs[0].font.bold = True

        regulatory_docs_paragraph = self._document.add_paragraph()
        regulatory_docs_paragraph.add_run("5. Нормативные документы:\n")
        for doc in self._file.regulatory_docs:
            regulatory_docs_paragraph.add_run(f"{doc}\n")
        regulatory_docs_paragraph.runs[0].font.bold = True

    def _add_summary_table(self) -> None:
        label_paragraph = self._document.add_paragraph()
        label_paragraph.add_run("6. Сводные данные:")
        label_paragraph.runs[0].font.bold = True

        table = self._document.add_table(rows=4, cols=2)
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(1.5)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.style = 'Table Grid'

        data = {
            "Уровень профессионального риска на рабочем месте по результатам оценки:": str(round(self._file.prof_risk, 2)) if self._file.prof_risk else '',
            "Показатель, учитывающий результаты специальной оценки условий труда:": str(self._file.k_factor) if self._file.k_factor is not None else '',
            "Итоговый уровень профессионального риска по результатам оценки:": str(round(self._file.result, 2)) if self._file.result else '',
            "По степени риска рабочее место отнесено к категории:": self._file.result_str if self._file.result_str else ''
        }

        for i, row in enumerate(table.rows):
            desc, val  = list(data.items())[i]

            cell_desc = row.cells[0]
            cell_desc.text = desc
            cell_desc.paragraphs[0].runs[0].font.size = Pt(10)
            cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cell_val = row.cells[1]
            cell_val.text = val
            cell_val.paragraphs[0].runs[0].font.size = Pt(10)
            cell_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if i == 3 and self._file.result_str:
                if self._file.result_str == "Высокий":
                    cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                elif self._file.result_str == "Средний":
                    cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(252, 186, 3)
                else:
                    cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        self._document.add_paragraph()

    def _add_risk_table(self) -> None:
        prof_risk_paragraph = self._document.add_paragraph()
        prof_risk_paragraph.add_run("7. Расчет уровней профессиональных рисков:")
        prof_risk_paragraph.runs[0].font.bold = True
        headers = [
            '№ п/п',
            'Номер опасности по перечню',
            'Опасность',
            'Опасное событие',
            'Качественное значение тяжести ущерба',
            'Баллы по тяжести ущерба',
            'Качественное значение подверженности опасности',
            'Баллы подверженности опасности',
            'Качественное значение вероятности возникновения опасности',
            'Вероятность возникновения опасности',
            'Сумма весовых коэффициентов',
            'Весовой коэффициент вероятности возникновения опасности',
            'Риски по идентифицированным опасностям',
            'Оценка значимости риска по отдельной опасности'
        ]
        table = self._document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        col_widths = [0.4, 0.6, 2.2, 2.2, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5]
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

        table.repeat_header_rows = True
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._protect_cell_from_splitting(hdr_cells[i])
        self._prevent_row_split(table.rows[0])

        weight_sum = sum(record.probability_pts for record in self._file.table) if self._file.table else 0

        for i, record in enumerate(self._file.table):
            row_cells = table.add_row().cells

            row_cells[0].text = str(i + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[1].text = str(record.n) if record.n else ''
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[2].text = record.danger if record.danger else ''

            row_cells[3].text = record.event if record.event else ''

            row_cells[4].text = record.damage if record.damage else ''
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[5].text = str(record.damage_pts) if record.damage_pts else ''
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[6].text = record.susceptibility if record.susceptibility else ''
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[7].text = str(record.susceptibility_pts) if record.susceptibility_pts else ''
            row_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[8].text = record.probability if record.probability else ''
            row_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[9].text = str(record.probability_pts) if record.probability_pts else ''
            row_cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if i == 0:
                row_cells[10].text = str(weight_sum)
            row_cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if record.probability_pts and weight_sum:
                k = round(record.probability_pts / weight_sum, 2)
                row_cells[11].text = str(k)

            else: row_cells[11].text = ''
            row_cells[11].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[12].text = str(record.identified_dangers_risks) if record.identified_dangers_risks else ''
            row_cells[12].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[13].text = record.rating if record.rating else ''
            row_cells[13].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

            self._prevent_row_split(table.rows[-1])
            for cell in row_cells:
                self._protect_cell_from_splitting(cell)

        table.rows[1].cells[10].merge(table.rows[-1].cells[10])
        self._document.add_paragraph()

    def _add_methods(self) -> None:
        methods_paragraph = self._document.add_paragraph()
        methods_paragraph.add_run("Общие меры по управлению рисками:\n")

        if not self._file.methods:
            methods_paragraph.add_run("")
            return

        for method in self._file.methods:
            methods_paragraph.add_run(f"–  {method}\n")
        methods_paragraph.runs[0].font.bold = True

    @staticmethod
    def _protect_cell_from_splitting(cell):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_lines_together = True

    @staticmethod
    def _prevent_row_split(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant_split = trPr.makeelement(qn('w:cantSplit'), {})
        trPr.append(cant_split)